"""
Document export service.

Pure functions — no DB access, no LLM calls.
All inputs must already be validated before calling these.

DOCX: python-docx  (installed)
PDF:  reportlab    (installed)
"""
from __future__ import annotations

import io
import re
import textwrap

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

# ── Markdown-lite parser ──────────────────────────────────────────────────────

def _parse_lines(text: str) -> list[tuple[str, str]]:
    """
    Returns list of (kind, content):
      heading1, heading2, heading3, bullet, blank, para
    """
    result: list[tuple[str, str]] = []
    for raw in text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            result.append(("blank", ""))
        elif line.startswith("### "):
            result.append(("heading3", line[4:].strip()))
        elif line.startswith("## "):
            result.append(("heading2", line[3:].strip()))
        elif line.startswith("# "):
            result.append(("heading1", line[2:].strip()))
        elif re.match(r"^[-*•]\s+", line):
            result.append(("bullet", re.sub(r"^[-*•]\s+", "", line).strip()))
        else:
            result.append(("para", line))
    return result


# ── DOCX helpers ──────────────────────────────────────────────────────────────

def _set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    # Also set theme font to avoid Word override
    style.element.rPr.rFonts.set(qn("w:asciiTheme"), "minorHAnsi")


def _add_heading(doc: Document, text: str, level: int) -> None:
    sizes = {1: 16, 2: 13, 3: 11}
    bold = {1: True, 2: True, 3: True}
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold.get(level, False)
    run.font.size = Pt(sizes.get(level, 11))
    if level == 1:
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x2D, 0x5B, 0xA3)
        p.runs[0].font.size = Pt(13)
    p.paragraph_format.space_before = Pt(8 if level > 1 else 12)
    p.paragraph_format.space_after = Pt(2)


def _add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text).font.size = Pt(10.5)
    p.paragraph_format.left_indent = Pt(14)
    p.paragraph_format.space_after = Pt(1)


def _add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(3)


def _build_docx(text: str, doc_title: str) -> bytes:
    doc = Document()
    _set_default_font(doc)

    # Narrow margins
    for section in doc.sections:
        section.top_margin = Pt(50)
        section.bottom_margin = Pt(50)
        section.left_margin = Pt(65)
        section.right_margin = Pt(65)

    # Document title (first line if heading1, else synthetic)
    lines = _parse_lines(text)

    for kind, content in lines:
        if kind == "blank":
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
        elif kind == "heading1":
            _add_heading(doc, content, 1)
        elif kind == "heading2":
            _add_heading(doc, content, 2)
        elif kind == "heading3":
            _add_heading(doc, content, 3)
        elif kind == "bullet":
            _add_bullet(doc, content)
        else:  # para
            if content.strip():
                _add_para(doc, content)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_cv_docx(cv_text: str) -> bytes:
    return _build_docx(cv_text, "CV")


def build_cover_letter_docx(letter_text: str) -> bytes:
    return _build_docx(letter_text, "Cover Letter")


# ── PDF helpers ───────────────────────────────────────────────────────────────

_BASE_STYLES = getSampleStyleSheet()

_STYLE_H1 = ParagraphStyle(
    "ExH1",
    parent=_BASE_STYLES["Normal"],
    fontSize=16,
    leading=20,
    textColor=(0x1E / 255, 0x3A / 255, 0x5F / 255),
    fontName="Helvetica-Bold",
    spaceAfter=4,
    spaceBefore=10,
)
_STYLE_H2 = ParagraphStyle(
    "ExH2",
    parent=_BASE_STYLES["Normal"],
    fontSize=13,
    leading=17,
    textColor=(0x2D / 255, 0x5B / 255, 0xA3 / 255),
    fontName="Helvetica-Bold",
    spaceAfter=3,
    spaceBefore=8,
)
_STYLE_H3 = ParagraphStyle(
    "ExH3",
    parent=_BASE_STYLES["Normal"],
    fontSize=11,
    leading=14,
    fontName="Helvetica-Bold",
    spaceAfter=2,
    spaceBefore=5,
)
_STYLE_BODY = ParagraphStyle(
    "ExBody",
    parent=_BASE_STYLES["Normal"],
    fontSize=10.5,
    leading=14,
    fontName="Helvetica",
    spaceAfter=4,
)
_STYLE_BULLET = ParagraphStyle(
    "ExBullet",
    parent=_BASE_STYLES["Normal"],
    fontSize=10.5,
    leading=14,
    fontName="Helvetica",
    leftIndent=14,
    bulletIndent=4,
    spaceAfter=2,
)

_STYLE_MAP = {
    "heading1": _STYLE_H1,
    "heading2": _STYLE_H2,
    "heading3": _STYLE_H3,
    "para":     _STYLE_BODY,
    "bullet":   _STYLE_BULLET,
}


def _safe_xml(text: str) -> str:
    """Escape characters that break reportlab XML parsing."""
    return (
        text
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _build_pdf(text: str) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=23 * mm,
        rightMargin=23 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )
    story: list = []

    for kind, content in _parse_lines(text):
        if kind == "blank":
            story.append(Spacer(1, 3 * mm))
        elif kind == "bullet":
            story.append(
                Paragraph(f"• {_safe_xml(content)}", _STYLE_BULLET)
            )
        elif kind in _STYLE_MAP:
            if content.strip():
                story.append(
                    Paragraph(_safe_xml(content), _STYLE_MAP[kind])
                )

    if not story:
        story.append(Paragraph("(empty)", _STYLE_BODY))

    doc.build(story)
    return buf.getvalue()


def build_cv_pdf(cv_text: str) -> bytes:
    return _build_pdf(cv_text)


def build_cover_letter_pdf(letter_text: str) -> bytes:
    return _build_pdf(letter_text)


# ── Copy-ready message generation ─────────────────────────────────────────────

def _extract_pitch(cover_letter: str, max_sentences: int = 3) -> str:
    """
    Pull the first substantive sentences from the cover letter for use
    in shorter outreach messages. Skips salutation lines.
    """
    skip_patterns = re.compile(
        r"^\s*(dear|hi|hello|to whom|re:|subject:|sincerely|regards|yours|best|kind)",
        re.IGNORECASE,
    )
    sentences: list[str] = []
    for line in cover_letter.splitlines():
        line = line.strip()
        if not line or skip_patterns.match(line):
            continue
        # Split line into sentences on '. ', '! ', '? '
        parts = re.split(r"(?<=[.!?])\s+", line)
        for part in parts:
            if part.strip():
                sentences.append(part.strip())
            if len(sentences) >= max_sentences:
                break
        if len(sentences) >= max_sentences:
            break
    return " ".join(sentences)


def build_hr_email(
    job_title: str,
    company_name: str,
    candidate_name: str,
    cover_letter: str,
) -> str:
    pitch = _extract_pitch(cover_letter, max_sentences=3)
    return textwrap.dedent(f"""\
        Subject: Application for {job_title} at {company_name}

        Dear Hiring Manager,

        I am writing to express my interest in the {job_title} position at {company_name}.

        {pitch}

        I have attached my CV and cover letter for your review. I would welcome the opportunity to discuss how my background aligns with your team's needs.

        Kind regards,
        {candidate_name}
    """).strip()


def build_linkedin_message(
    job_title: str,
    company_name: str,
    candidate_name: str,
    cover_letter: str,
) -> str:
    pitch = _extract_pitch(cover_letter, max_sentences=2)
    return textwrap.dedent(f"""\
        Hi,

        I came across the {job_title} role at {company_name} and wanted to reach out directly. {pitch}

        Would you be open to a quick conversation? Happy to share more about my experience.

        Best,
        {candidate_name}
    """).strip()
